from __future__ import annotations
import os
from typing import List, Optional
from fastapi import FastAPI, UploadFile, File, Depends, HTTPException, Header, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse
from pydantic import BaseModel
from docx import Document as DocxDocument
from .config import settings
from .auth import authenticate_user, create_access_token, decode_token
from .ingest import build_faiss_from_dir, SUPPORTED_EXTS, CHUNK_SIZE, CHUNK_OVERLAP, extract_text_from_file
from .qa import answer_question
from langchain.text_splitter import RecursiveCharacterTextSplitter
import mammoth
import pandas as pd
import html
import re

app = FastAPI(title="DocChat API", version="0.2.0")

origins = [o.strip() for o in settings.CORS_ORIGINS.split(",")]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class LoginRequest(BaseModel):
    username: str
    password: str

class LoginResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"
    role: str

def get_current_user(authorization: Optional[str] = Header(None)):
    if not authorization or not authorization.lower().startswith("bearer "):
        raise HTTPException(status_code=401, detail="Missing bearer token")
    token = authorization.split()[1]
    payload = decode_token(token)
    if not payload:
        raise HTTPException(status_code=401, detail="Invalid token")
    return payload

@app.get("/health")
def health():
    return {"status": "ok"}

@app.post("/auth/login", response_model=LoginResponse)
def login(req: LoginRequest):
    user = authenticate_user(req.username, req.password)
    if not user:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    token = create_access_token({"sub": user["username"], "role": user["role"]})
    return {"access_token": token, "role": user["role"]}

@app.post("/ingest/upload")
async def upload(files: List[UploadFile] = File(...), user=Depends(get_current_user)):
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin only")
    saved = []
    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
    for f in files:
        ext = os.path.splitext(f.filename)[1].lower()
        if ext not in SUPPORTED_EXTS:
            continue
        path = os.path.join(settings.UPLOAD_DIR, f.filename)
        with open(path, "wb") as out:
            out.write(await f.read())
        saved.append(f.filename)
    return {"saved": saved, "message": "Uploaded."}

@app.post("/ingest/build")
def ingest_build(user=Depends(get_current_user)):
    if user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin only")
    ok = build_faiss_from_dir(settings.UPLOAD_DIR, settings.INDEX_DIR)
    return {"indexed": ok}

class ChatRequest(BaseModel):
    question: str

@app.post("/chat")
def chat(req: ChatRequest, user=Depends(get_current_user)):
    try:
        res = answer_question(req.question)
        os.makedirs(settings.INDEX_DIR, exist_ok=True)
        with open(os.path.join(settings.INDEX_DIR, "latest_answer.txt"), "w") as f:
            f.write(res.get("answer", ""))
        return res
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/export/summary.docx")
def export_summary(user=Depends(get_current_user)):
    latest_path = os.path.join(settings.INDEX_DIR, "latest_answer.txt")
    if not os.path.exists(latest_path):
        raise HTTPException(status_code=400, detail="No summary to export yet.")
    with open(latest_path, "r") as f:
        text = f.read().strip()
    doc = DocxDocument()
    doc.add_heading("DocChat Summary", level=1)
    doc.add_paragraph(text or "(empty)")
    out_path = os.path.join(settings.INDEX_DIR, "summary.docx")
    doc.save(out_path)
    return FileResponse(out_path, filename="summary.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.get("/view", response_class=HTMLResponse)
def view(source: str = Query(...), chunk: int = Query(...)):
    """Return formatted HTML preview of a specific file with the cited chunk highlighted."""
    path = os.path.join(settings.UPLOAD_DIR, source)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail="Source file not found.")
    # Recreate the same split to find the chunk text
    raw_text = extract_text_from_file(path)
    splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    chunks = splitter.split_text(raw_text)
    if chunk < 0 or chunk >= len(chunks):
        raise HTTPException(status_code=400, detail="Chunk index out of range.")
    target = chunks[chunk][:300]  # first 300 chars for matching
    
    ext = os.path.splitext(source)[1].lower()
    html_body = ""
    try:
        if ext == ".docx":
            with open(path, "rb") as f:
                result = mammoth.convert_to_html(f)
                html_body = result.value  # HTML string
        elif ext == ".xlsx":
            xls = pd.read_excel(path, sheet_name=None)
            parts = []
            for sheet, df in xls.items():
                parts.append(f"<h3>Sheet: {html.escape(sheet)}</h3>" + df.to_html(index=False, escape=False))
            html_body = "".join(parts)
        elif ext == ".csv":
            import pandas as pd
            df = pd.read_csv(path)
            html_body = df.to_html(index=False, escape=False)
        else:
            html_body = f"<pre>{html.escape(raw_text)}</pre>"
    except Exception as e:
        html_body = f"<pre>{html.escape(raw_text)}</pre>"

    # Highlight the first occurrence of the target text


    safe_target = re.escape(target.strip())
    pattern = re.compile(safe_target, re.IGNORECASE)
    def add_anchor_once(match):
        return f"<a id='chunk-{chunk}'></a><mark style='background:#e6f4ea;border-radius:4px;'>{match.group(0)}</mark>"
    
    html_body, n = pattern.subn(add_anchor_once, html_body, count=1)
    if n == 0:
        # fallback if not found, add anchor to top
        html_body = f"<a id='chunk-{chunk}'></a>" + html_body
    

    # Wrap in minimal page with TD-like styles
    page = f"""
    <html>
      <head>
        <meta charset='utf-8'/>
        <style>
          body {{ font-family: Helvetica, Arial, sans-serif; background: #fff; color: #001f1f; }}
          h3 {{ color: #007c41; }}
          table {{ border-collapse: collapse; width: 100%; }}
          th, td {{ border: 1px solid #e6e6e6; padding: 6px 8px; text-align: left; }}
          mark {{ background:#e6f4ea; }}
        </style>
      </head>
      <body>
        <h2 style="margin-top:0">{html.escape(source)}</h2>
        {html_body}
      </body>
    </html>
    """
    return HTMLResponse(page)
